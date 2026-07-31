"""Convierte el informe Quarto aprobado en un documento Word editable."""

from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "report" / "informe.qmd"
TARGET = ROOT / "outputs" / "reports" / "informe_final.docx"

lines = SOURCE.read_text(encoding="utf-8").splitlines()
document = Document()
document.core_properties.title = "Análisis del contexto nacional y global"
document.core_properties.subject = "Transformación digital y empleo"
style = document.styles["Normal"]
style.font.name = "Aptos"
style.font.size = Pt(11)
section = document.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header.paragraphs[0].text = "Universidad Técnica de Cotopaxi · Carrera de Economía"
section.footer.paragraphs[0].alignment = 1
section.footer.paragraphs[0].add_run("Informe de trabajo · datos 2020–2024")
document.add_heading("Análisis del contexto nacional y global", 0)
document.add_paragraph("Transformación digital y empleo", style="Subtitle")
document.add_paragraph("Ecuador, Argentina, Colombia y Venezuela · 2020–2024")
document.add_paragraph("Universidad Técnica de Cotopaxi · Carrera de Economía")
document.add_paragraph("Documento editable; resultados exploratorios, no causales.")
document.add_page_break()

in_yaml = False
index = 0
while index < len(lines):
    line = lines[index].strip()
    if line == "---":
        in_yaml = not in_yaml
        index += 1
        continue
    if in_yaml or not line:
        index += 1
        continue
    if line.startswith("#"):
        level = len(line) - len(line.lstrip("#"))
        document.add_heading(line[level:].strip(), level=min(level, 3))
        index += 1
        continue
    if line.startswith("|"):
        table_lines = []
        while index < len(lines) and lines[index].strip().startswith("|"):
            if not set(lines[index].strip()) <= {"|", "-", ":", " "}:
                table_lines.append(lines[index].strip())
            index += 1
        cells = [[cell.strip() for cell in row.strip("|").split("|")] for row in table_lines]
        if cells:
            table = document.add_table(rows=1, cols=len(cells[0]))
            table.style = "Table Grid"
            for cell, text in zip(table.rows[0].cells, cells[0]):
                cell.text = text
            for row in cells[1:]:
                new_cells = table.add_row().cells
                for cell, text in zip(new_cells, row):
                    cell.text = text
        continue
    if line.startswith("- "):
        document.add_paragraph(line[2:], style="List Bullet")
    else:
        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", line)
        text = text.replace("`", "").replace("$", "")
        document.add_paragraph(text)
    index += 1

TARGET.parent.mkdir(parents=True, exist_ok=True)
document.save(TARGET)
print(f"Documento Word generado: {TARGET}")
